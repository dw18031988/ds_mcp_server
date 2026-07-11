#!/usr/bin/env python3
from __future__ import annotations
import argparse, datetime, hashlib, json, re, shutil, zipfile
from pathlib import Path
from typing import Iterable

try:
    import yaml
except Exception as exc:
    raise SystemExit("PyYAML is required: pip install pyyaml") from exc

CANONICAL = "Coding_Project_Governance_v1.0.md"
DOCX_NAME = "Coding_Project_Governance_v1.0.docx"
MANIFEST_NAME = "policy_manifest.v1.0.json"
BUNDLE_NAME = "Generic_Coding_Governance_v1.0_Bundle.zip"

CRITICAL_IDS = [
    "TRUST-01","TRUST-02","TRUST-03","TRUST-04","TRUST-06",
    "GATE-01","GATE-02","GATE-03","GATE-04","GATE-05",
    "PROFILE-01","PROFILE-02","PROFILE-03","PROFILE-04",
    "PRE-02","PRE-03","PRE-04","PRE-05",
    "GH-02","GH-03","CMD-01","CMD-02","CMD-03","CMD-04","CMD-05",
    "DATA-01","DATA-02","DATA-03","SEC-01","SEC-02",
    "VAL-04","VAL-05","PR-02","CI-01","CI-02","CI-03","CI-04","CI-05",
    "POLICY-01","POLICY-02","POLICY-03","POLICY-04","POLICY-05","BAN-01"
]

AUTHORITATIVE_FILES = [CANONICAL]
DERIVED_FILES = [
    "Coding_Project_Governance_v1.0_Compact.md",
    DOCX_NAME,
    "ChatGPT_Project_Bootstrap.template.md",
    "Project_Profile.template.yaml",
    "Project_Extension.template.md",
    "policy_machine.v1.0.json",
    "Enforcement_Checklist.md",
    "Lessons_Learned.md",
    "Migration_Guide_DS_MCP_to_Generic.md",
    "RedTeam_ReReview.md",
]
SUPPORT_FILES = [
    "approval_envelope.schema.json",
    "change_plan.schema.json",
    "project_profile.schema.json",
    "project_extension.schema.json",
    "policy_manifest.schema.json",
    "enforce_coding_policy_v1_0.py",
    "build_generic_policy.py",
]
PROJECT_FILES = [
    "projects/rental-home/project-profile.yaml",
    "projects/rental-home/project-extension.md",
    "projects/rental-home/chatgpt-bootstrap.md",
    "projects/ds-mcp/project-profile.yaml",
    "projects/ds-mcp/project-extension.md",
    "projects/ds-mcp/chatgpt-bootstrap.md",
    "projects/pm-skills/project-profile.yaml",
    "projects/pm-skills/project-extension.md",
    "projects/pm-skills/chatgpt-bootstrap.md",
]

def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()

def utc_now() -> str:
    return datetime.datetime.now(datetime.timezone.utc).replace(microsecond=0).isoformat().replace("+00:00","Z")

def update_front_matter(path: Path, core_hash: str, generated_at: str) -> None:
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        return
    end = text.find("\n---\n",4)
    if end < 0:
        return
    meta = yaml.safe_load(text[4:end]) or {}
    if meta.get("authoritative") is False:
        if "canonical_sha256" in meta:
            meta["canonical_sha256"] = core_hash
        if "core_sha256" in meta:
            meta["core_sha256"] = core_hash
        if "generated_at" in meta:
            meta["generated_at"] = generated_at
    meta = {k:v for k,v in meta.items() if v is not None}
    path.write_text("---\n"+yaml.safe_dump(meta,sort_keys=False).strip()+"\n---\n"+text[end+5:],encoding="utf-8")

def replace_hash_text(path: Path, old_pattern: str, core_hash: str) -> None:
    text = path.read_text(encoding="utf-8")
    text = re.sub(old_pattern, core_hash, text)
    path.write_text(text,encoding="utf-8")

def set_run_font(run, name: str, size=None, bold=None, color=None):
    from docx.oxml.ns import qn
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def mark_header_row(row):
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val","true")
    trPr.append(tblHeader)

def shade_cell(cell, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def add_code(doc, text: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    pf = p.paragraph_format
    pf.left_indent = Pt(12)
    pf.right_indent = Pt(6)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run,"Liberation Mono",Pt(8.5))
    return p

def add_markdown(doc, text: str):
    from docx.shared import Pt
    in_code=False
    code_lines=[]
    for raw in text.splitlines():
        line=raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code(doc,"\n".join(code_lines))
                code_lines=[]
                in_code=False
            else:
                in_code=True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(),level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(),level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(),level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(),level=3)
        elif re.match(r"^\d+\.\s",line):
            p=doc.add_paragraph()
            p.paragraph_format.left_indent=Pt(18)
            p.paragraph_format.first_line_indent=Pt(-12)
            p.add_run(line)
        elif line.startswith("- "):
            p=doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        elif line.startswith("|") and line.endswith("|"):
            # Tables are handled only for compact standalone blocks; leave as monospace to avoid malformed parsing.
            add_code(doc,line)
        elif line.startswith("> "):
            p=doc.add_paragraph()
            p.paragraph_format.left_indent=Pt(18)
            r=p.add_run(line[2:])
            r.italic=True
        else:
            p=doc.add_paragraph()
            p.add_run(line)
    if code_lines:
        add_code(doc,"\n".join(code_lines))

def build_docx(root: Path, core_hash: str):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn

    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Inches(0.65); sec.bottom_margin=Inches(0.65)
    sec.left_margin=Inches(0.72); sec.right_margin=Inches(0.72)

    styles=doc.styles
    styles["Normal"].font.name="Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"),"Aptos")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"),"Aptos")
    styles["Normal"].font.size=Pt(9.5)
    for name,size,color in [("Title",28,"0F172A"),("Heading 1",17,"0F3D73"),("Heading 2",13,"1F5C99"),("Heading 3",11,"334155")]:
        st=styles[name]
        st.font.name="Aptos Display" if name=="Title" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:ascii"),st.font.name)
        st._element.rPr.rFonts.set(qn("w:hAnsi"),st.font.name)
        st.font.size=Pt(size)
        st.font.color.rgb=RGBColor.from_string(color)
        st.font.bold=True

    title=doc.add_paragraph()
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    title.space_after=Pt(8)
    r=title.add_run("Generic Coding Project Governance")
    set_run_font(r,"Aptos Display",Pt(28),True,RGBColor(15,23,42))
    sub=doc.add_paragraph()
    sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sub.add_run("Canonical Code of Conduct + Project Profile Model")
    set_run_font(r,"Aptos",Pt(14),False,RGBColor(71,85,105))
    meta=doc.add_paragraph()
    meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=meta.add_run(f"Version 1.0  |  Canonical SHA-256: {core_hash}")
    set_run_font(r,"Liberation Mono",Pt(8),False,RGBColor(71,85,105))
    doc.add_paragraph()
    note=doc.add_paragraph()
    note.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=note.add_run("Non-authoritative human-readable derivative. The Markdown canonical source controls.")
    r.italic=True
    set_run_font(r,"Aptos",Pt(10),False,RGBColor(153,27,27))

    doc.add_page_break()
    doc.add_heading("How to use this package",level=1)
    p=doc.add_paragraph("Load the canonical core once. For each task, load exactly one Project Profile and its Extensions. The code of conduct remains identical; only repository and project context vary.")
    table=doc.add_table(rows=1,cols=3)
    table.style="Table Grid"
    headers=["Layer","Purpose","Authority"]
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.text=h; shade_cell(cell,"DCEAF7")
    mark_header_row(table.rows[0])
    for row in [
        ("Canonical core","Trust, approvals, Git, command, security, validation, PR/CI","Authoritative"),
        ("Project Profile","Repository, connector, context, commands, CI, providers","Context only"),
        ("Project Extension","Stricter project conventions","Tighten only"),
    ]:
        cells=table.add_row().cells
        for i,val in enumerate(row): cells[i].text=val

    doc.add_heading("Included Project Profiles",level=1)
    ptable=doc.add_table(rows=1,cols=5); ptable.style="Table Grid"
    for i,h in enumerate(["Project","Repository","Connector","Write status","Extension"]):
        ptable.rows[0].cells[i].text=h; shade_cell(ptable.rows[0].cells[i],"DCEAF7")
    mark_header_row(ptable.rows[0])
    for rel in PROJECT_FILES:
        if not rel.endswith("project-profile.yaml"): continue
        prof=yaml.safe_load((root/rel).read_text(encoding="utf-8"))
        repo=prof["repository"]
        cells=ptable.add_row().cells
        vals=[prof["project"]["name"],f"{repo['owner']}/{repo['name']}",prof["connectors"]["git_provider"],
              "Enabled" if repo["write_enabled"] else "Locked",", ".join(prof["extensions"])]
        for i,v in enumerate(vals): cells[i].text=str(v)

    doc.add_page_break()
    core=(root/CANONICAL).read_text(encoding="utf-8")
    # Strip front matter from DOCX body.
    if core.startswith("---\n"):
        end=core.find("\n---\n",4)
        if end>=0: core=core[end+5:]
    add_markdown(doc,core)

    # Header/footer.
    for section in doc.sections:
        hp=section.header.paragraphs[0]
        hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        rr=hp.add_run("Generic Coding Governance v1.0")
        set_run_font(rr,"Aptos",Pt(8),False,RGBColor(100,116,139))
        fp=section.footer.paragraphs[0]
        fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=fp.add_run("Canonical source: Coding_Project_Governance_v1.0.md")
        set_run_font(rr,"Aptos",Pt(8),False,RGBColor(100,116,139))

    props=doc.core_properties
    props.title="Generic Coding Project Governance v1.0"
    props.subject="Cross-project coding governance"
    props.author="Generic Coding Governance Generator"
    props.last_modified_by="Generic Coding Governance Generator"
    props.comments="Generated from the canonical Markdown source."
    props.created=datetime.datetime.now(datetime.timezone.utc)
    props.modified=datetime.datetime.now(datetime.timezone.utc)
    doc.save(root/DOCX_NAME)

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument("--dir",default=".")
    args=ap.parse_args()
    root=Path(args.dir)
    core_hash=sha256(root/CANONICAL)
    generated_at=utc_now()

    # Update non-authoritative Markdown metadata and visible bootstrap hashes.
    for rel in DERIVED_FILES+PROJECT_FILES:
        path=root/rel
        if not path.exists() or path.suffix.lower() not in {".md",".yaml",".yml",".json"}:
            continue
        if path.suffix.lower()==".md":
            update_front_matter(path,core_hash,generated_at)
            text=path.read_text(encoding="utf-8")
            text=re.sub(r"SHA-256: `[a-f0-9]{64}`",f"SHA-256: `{core_hash}`",text)
            text=re.sub(r"core SHA-256: `[a-f0-9]{64}`",f"core SHA-256: `{core_hash}`",text)
            text=re.sub(r"Canonical generic SHA: `[a-f0-9]{64}`",f"Canonical generic SHA: `{core_hash}`",text)
            path.write_text(text,encoding="utf-8")
        elif path.name=="policy_machine.v1.0.json":
            data=json.loads(path.read_text(encoding="utf-8"))
            data["canonical_sha256"]=core_hash; data["generated_at"]=generated_at
            path.write_text(json.dumps(data,indent=2)+"\n",encoding="utf-8")
        elif path.suffix.lower() in {".yaml",".yml"} and "project-profile" not in path.name:
            pass

    build_docx(root,core_hash)

    # Manifest excludes itself, bundle, and transient QA.
    hash_files=AUTHORITATIVE_FILES+DERIVED_FILES+SUPPORT_FILES+PROJECT_FILES
    file_hash={rel:sha256(root/rel) for rel in hash_files}
    manifest={
        "version":"1.0","canonical_file":CANONICAL,"canonical_sha256":core_hash,
        "generated_at":generated_at,"generator":"build_generic_policy.py@1.0",
        "authoritative_files":AUTHORITATIVE_FILES,"derived_files":DERIVED_FILES,
        "support_files":SUPPORT_FILES,"project_profiles":[x for x in PROJECT_FILES if x.endswith("project-profile.yaml")],
        "critical_rule_ids":CRITICAL_IDS,"file_sha256":file_hash,
    }
    (root/MANIFEST_NAME).write_text(json.dumps(manifest,indent=2)+"\n",encoding="utf-8")

    bundle=root/BUNDLE_NAME
    if bundle.exists(): bundle.unlink()
    include=hash_files+[MANIFEST_NAME]
    with zipfile.ZipFile(bundle,"w",compression=zipfile.ZIP_DEFLATED) as zf:
        for rel in include:
            zf.write(root/rel,arcname=rel)
    print(f"canonical_sha256={core_hash}")
    print(f"docx={root/DOCX_NAME}")
    print(f"manifest={root/MANIFEST_NAME}")
    print(f"bundle={bundle}")

if __name__=="__main__":
    main()
